import os
from typing import List, Dict
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
from pathlib import Path
from pdf2image import convert_from_path
from bs4 import BeautifulSoup
import json

class TelecomEgyptDocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.htm', 
                                 '.png', '.jpg', '.jpeg', '.tiff', '.bmp']

    def process_pdf(self, file_path: str) -> str:

        text = []
        try:
            # Step 1: Try normal text extraction
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text + "\n")

            # Step 2: If no text found, use OCR
            if not text:
                images = convert_from_path(file_path)
                for img in images:
                    ocr_text = pytesseract.image_to_string(img)
                    text.append(ocr_text + "\n")

            return text

        except Exception as e:
            print(f"Error processing PDF {file_path}: {str(e)}")
            return ""

    def process_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return text

        except Exception as e:
            print(f"Error processing DOCX {file_path}: {str(e)}")
            return ""


    def process_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            print(f"Error processing TXT {file_path}: {str(e)}")
            return ""
    
    def process_html(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')

                for element in soup(["script", "style", "noscript"]):
                    element.decompose()

            return soup.get_text(separator=' ', strip=True)

        except Exception as e:
            print(f"Error processing HTML {file_path}: {str(e)}")
            return ""

    def process_image(self, file_path: str) -> str:
        try:
            image = Image.open(file_path)
            # Try to extract Arabic and English text
            text = pytesseract.image_to_string(image, lang='ara+eng')
            return [str(text)]
        except Exception as e:
            print(f"Error processing image {file_path}: {str(e)}")
            return ""

    def process_document(self, file_path: str) -> Dict:
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Extract text based on file type
        if file_ext == '.pdf':
            content = self.process_pdf(str(file_path))
        elif file_ext == '.docx':
            content = self.process_docx(str(file_path))
        elif file_ext in ['.txt']:
            content = self.process_txt(str(file_path))
        elif file_ext in ['.html', '.htm']:
            content = self.process_html(str(file_path))
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            content = self.process_image(str(file_path))
        else:
            content = ""
        
        return {
            'filename': file_path.name,
            'filepath': str(file_path),
            'file_type': file_ext,
            'content': content,
            'content_length': len(content),
            'source': 'user_upload'
        }
    

    def save_to_json(self, data: Dict, filename: str):
        try:
            existing_data = []
            if os.path.exists(filename):
                try:
                    with open(filename, 'r', encoding='utf-8') as f:
                        content = f.read()
                        if content.strip():
                            existing_data = json.loads(content)
                            if isinstance(existing_data, dict):
                                existing_data = [existing_data]
                except json.JSONDecodeError:
                    pass  # Start with empty list if file is invalid JSON

            existing_data.append(data)

            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
            print(f"Data successfully saved to {filename}")
        except Exception as e:
             print(f"Error saving data to JSON: {str(e)}")

    def process_multiple_documents(self, file_paths: List[str], output_filename: str):
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                content_list=result['content']
                for page_number,page_text in enumerate(content_list):
                    if page_text:
                        result['content']=page_text
                        result['content_length']=len(page_text)
                        result['page_number']=page_number + 1
                        self.save_to_json(result, output_filename)
                    else:
                        print(f"✗ Empty content: {file_path}")
            except Exception as e:
                print(f"✗ Failed: {file_path} - {str(e)}")
        
        return  


def docs_main(file_paths: List[str], output_path: str = None):
    processor = TelecomEgyptDocumentProcessor()
    processor.process_multiple_documents(file_paths, output_path)
    return 

